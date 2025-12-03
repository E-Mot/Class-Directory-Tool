from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches
import os, re

#------------------------------------------- Get the path to the desktop -------------------------------------------

DESKTOP_PATH = os.path.join(os.path.expanduser("~"), "Desktop")

#------------------------------------------- Creating the 'Class template' folder -------------------------------------------

def create_outer_folder():
    # Define the name of the new folder
    outer_folder = "Class template"
    # Combine the desktop path and folder name
    folder_path = os.path.join(DESKTOP_PATH, outer_folder)
    # Create the folder and allow re-creation for experimentation
    os.makedirs(folder_path, exist_ok = True)

#------------------------------------------- Creating the first level of sub-directories with the 'Class template' folder -------------------------------------------

def create_sub_dirs():
    sub_folders = ["1. Units", "2. Class materials", "3. Extra resources"]
    sub_directory_path = os.path.join(DESKTOP_PATH, "Class template")
    for folder_name in sub_folders:
        sub_folder_path = os.path.join(sub_directory_path, folder_name)
        os.makedirs(sub_folder_path, exist_ok = True)



#------------------------------------------- Creating sub-directories within the '1. Units' folder -------------------------------------------

def create_units_sub_dirs():
     sub_folders = [".Unit 1", ".Unit 2", ".Unit 3", ".Unit 4", ".Unit 5", ".Unit 6", ".Unit 7", ".Unit 8", ".Unit 9", ".Unit 10", "Bonus material"]
     sub_directory_path = os.path.join(DESKTOP_PATH, "Class template", "1. Units")
     for folder_name in sub_folders:
         sub_folder_path = os.path.join(sub_directory_path, folder_name)
         os.makedirs(sub_folder_path, exist_ok = True)



#------------------------------------------- Creating sub-directories within each '.Unit X' folder -------------------------------------------

def create_unit_folder_sub_dirs():
     sub_folders = ["Activities, labs, quizzes", "Assignments", "Discussions", "Reading & Notes", "Seminar"]
     sub_directory_path = os.path.join (DESKTOP_PATH, "Class template", "1. Units")
     for unit_name in os.listdir(sub_directory_path):
        if unit_name == "Bonus material":
            continue
        else:
            unit_path = os.path.join(sub_directory_path, unit_name)
            for sub in sub_folders:
                dir_path = os.path.join(unit_path, sub)
                os.makedirs(dir_path, exist_ok = True)

def create_activities_sub_dirs():
    sub_folders = ["Activities", "Labs", "Quizzes"]
    sub_folder_path = os.path.join(DESKTOP_PATH, "Class template", "1. Units")
    for folder_name in os.listdir(sub_folder_path):
        if folder_name == "Bonus material":
            continue
        else:
            folder_path = os.path.join(sub_folder_path, folder_name)
            activities_path = os.path.join(folder_path, "Activities, labs, quizzes")
            for sub in sub_folders:
                dir_path = os.path.join(activities_path, sub)
                os.makedirs(dir_path, exist_ok = True)

def create_reading_sub_dirs():
    sub_folders = [".Notes", "Articles", "Book(s)", "Unit documents", "Videos"]
    sub_folder_path = os.path.join(DESKTOP_PATH,"Class template", "1. Units")
    for folder_name in os.listdir(sub_folder_path):
        if folder_name == "Bonus material":
            continue
        else:
            folder_path = os.path.join(sub_folder_path, folder_name)
            reading_path = os.path.join(folder_path, "Reading & Notes")
            for sub in sub_folders:
                dir_path = os.path.join(reading_path, sub)
                os.makedirs(dir_path, exist_ok = True)

#------------------------------------------- Creating files within each 'Discussions' and 'Seminar' folders -------------------------------------------

def create_discussion_file():
    sub_folder_path = os.path.join(DESKTOP_PATH, "Class template", "1. Units")
    for folder_name in os.listdir(sub_folder_path):
        folder_path = os.path.join(sub_folder_path, folder_name)
        for sub in os.listdir(folder_path):
            if sub == "Discussions":
                file_path = os.path.join(folder_path, "Discussions", "Discussion board posts.docx")
                doc = Document()

                sections = doc.sections
                for section in sections:
                    section.top_margin = Inches(1)
                    section.bottom_margin = Inches(1)
                    section.left_margin = Inches(1)
                    section.right_margin = Inches(1)

                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run("Class #")
                run.bold = True
                font = run.font
                font.name = 'Times New Roman'
                font.size = Pt(20)

                p = paragraph._p
                pPr = p.get_or_add_pPr()
                pbdr = OxmlElement('w:pBdr')
                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), '12')
                bottom.set(qn('w:space'), '1')
                bottom.set(qn('w:color'), '000000')
                pbdr.append(bottom)
                pPr.append(pbdr)

                paragraph_2 = doc.add_paragraph()
                run_2 = paragraph_2.add_run("xx/xx/xx")
                run_2.bold = True
                run_2.underline = True
                font_2 = run_2.font
                font_2.name = 'Times New Roman'
                font_2.size = Pt(12)

                doc.save(file_path)
 

def create_seminar_file():
    file_list = ["Unit 1 - Seminar - Chat log", "Unit 2 - Seminar - Chat log", "Unit 3 - Seminar - Chat log", "Unit 4 - Seminar - Chat log", "Unit 5 - Seminar - Chat log", "Unit 6 - Seminar - Chat log", "Unit 7 - Seminar - Chat log", "Unit 8 - Seminar - Chat log", "Unit 9 - Seminar - Chat log", "Unit 10 - Seminar - Chat log"]
    file_index = 0
    sub_folder_path = os.path.join(DESKTOP_PATH, "Class template", "1. Units")

    def extract_unit_number(folder_name):
        match = re.search(r"\d+", folder_name)  # looks for a number in the directory string using a regex pattern
        if match:                               # if a match is found...
            return int(match.group())           # convert it to an integer and return it to the sorted() function for sorting order
        else:                                   # if no digits are found...
            return 0                            # return 0 as a fallback
    
    for folder_name in sorted(os.listdir(sub_folder_path), key = extract_unit_number): # the sorted() function calls the extract_unit_number() function via the 'key' parameter and passes in the os.listdir() directory string
        folder_path = os.path.join(sub_folder_path, folder_name)
        for sub in os.listdir(folder_path):
            if sub == "Seminar":
                file_path = os.path.join(folder_path, "Seminar", f"{file_list[file_index]}.docx")

                doc = Document()

                sections = doc.sections
                for section in sections:
                    section.top_margin = Inches(1)
                    section.bottom_margin = Inches(1)
                    section.left_margin = Inches(1)
                    section.right_margin = Inches(1)

                paragraph = doc.add_paragraph()
                run = paragraph.add_run("Disclaimer:")
                run.underline = True
                font = run.font
                font.name = 'Helvetica'
                font.size = Pt(10.5)

                run.add_break()

                paragraph_2 = doc.add_paragraph()
                run_2 = paragraph_2.add_run("The chat log recorded on the school portal is not the full chat log from the zoom seminar. As a result, some posts from the chat log are not shown here.")
                font_2 = run_2.font
                font_2.name = 'Helvetica'
                font_2.size = Pt(10.5)

                run_2.add_break()

                p = paragraph_2._p
                pPr = p.get_or_add_pPr()
                pbdr = OxmlElement('w:pBdr')
                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), '12')
                bottom.set(qn('w:space'), '1')
                bottom.set(qn('w:color'), '000000')
                pbdr.append(bottom)
                pPr.append(pbdr)

                doc.save(file_path)

                if file_index < 9:
                    file_index += 1
                else:
                    break
                

                    


#------------------------------------------- Creating sub-directories within the '2. Class materials' folder -------------------------------------------

def create_class_materials_sub_dirs():
    sub_folders = ["1. Start here", "2. Course syllabus", "3. Reading", "4. Course resources", "5. Academic tools", "6. Virtual office", "7. Unit 1", "8. Unit 2", "9. Unit 3", "10. Unit 4", "11. Unit 5", "12. Unit 6", "13. Unit 7", "14. Unit 8", "15. Unit 9", "16. Unit 10"]
    sub_directory_path = os.path.join(DESKTOP_PATH, "Class template", "2. Class materials")
    for folder_name in sub_folders:
        sub_folder_path = os.path.join(sub_directory_path, folder_name)
        os.makedirs(sub_folder_path, exist_ok = True)

def create_resources_sub_dirs():
    sub_folders = ["Grading rubrics", "Links and documents"]
    sub_folder_path = os.path.join(DESKTOP_PATH, "Class template", "2. Class materials", "4. Course resources")
    for folder_name in sub_folders:
        dir_path = os.path.join(sub_folder_path, folder_name)
        os.makedirs(dir_path, exist_ok = True)

def create_rubrics_sub_dirs():
    sub_folders = ["Assignments", "Discussions", "Seminars"]
    sub_folder_path = os.path.join(DESKTOP_PATH, "Class template", "2. Class materials", "4. Course resources", "Grading rubrics")
    for folder_name in sub_folders:
        dir_path = os.path.join(sub_folder_path, folder_name)
        os.makedirs(dir_path, exist_ok = True)

#------------------------------------------- Creating files within the '6. Virtual office' folder -------------------------------------------

def create_office_file():
    file_path = os.path.join(DESKTOP_PATH, "Class template", "2. Class materials", "6. Virtual office", "Note.docx")
    doc = Document()

    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    paragraph = doc.add_paragraph()
    run = paragraph.add_run("This section contained a link to the classroom virtual office.")
    run.bold = True
    font = run.font
    font.name = 'Times New Roman'
    font.size = Pt(14)

    doc.save(file_path)

#------------------------------------------- Executing the functions -------------------------------------------

def execute_template_functions():
    create_outer_folder()
    create_sub_dirs()
    create_units_sub_dirs()
    create_unit_folder_sub_dirs()
    create_activities_sub_dirs()
    create_reading_sub_dirs()
    create_class_materials_sub_dirs()
    create_resources_sub_dirs()
    create_rubrics_sub_dirs()
    create_discussion_file()
    create_seminar_file()
    create_office_file()